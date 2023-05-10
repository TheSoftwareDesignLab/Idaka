# -*- coding: utf-8 -*-
"""
Created on Mon Oct 10 08:41:59 2022

@author: Laura
"""
"BM25 --https://www.quora.com/How-does-BM25-work"

from math import log
import collections
""" import matplotlib.pyplot as plt """
import re
import numpy as np
import string
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
""" import docx """
import json

""" import openpyxl """
""" from Query_expansion import expansion """
sw=stopwords.words("english")
def BM25(q,docs,vectorizer, original,b,k1):

    k2 = 0
    k3=1000

    i=0
    j=0
    result = dict()
    avgdl=average_dl(docs)
    for doc in docs:
        result[i]=BM25_doc(q,doc,avgdl,k1,k2,k3,b,vectorizer)
        i+=1
    result= sorted(result.items(), key=lambda x: x[1], reverse=True)
    ans={"practices":[],"scores":[]}
    for r in result:
        if r[1]>0:
            ans["practices"].append(r[0])
            ans["scores"].append(r[1])
            # print(original[r[0]])
            # # print(docs[r[0]])
            # print("Similarity:")
            # print(r[1])
    return json.dumps(ans)
    
    
def BM25_doc(q,doc, avgdl,k1,k2,k3,b,vectorizer):
    ans=0
    query=q.split(' ')
    dl=len(doc.split(' '))
    K=calculateK(dl, avgdl,b,k1)
    for term in query:
        ans+=(idf(term,vectorizer)*(((k1+1)*tf(term,doc))/(K+tf(term,doc)))*((k3+1)*tf(term,q))/(k3+tf(term,q))+k2*len(query)*((avgdl-dl)/(avgdl+dl)))   
    return ans                                                                    
def calculateK(dl,avgdl,b,k1):
    return k1*((1-b)+(b*dl/avgdl))

def idf(term,vectorizer):
    if vectorizer.vocabulary_.get(term)!= None:
        return vectorizer.idf_[vectorizer.vocabulary_.get(term)]
    else:
        return 0

def tf(term, sentence):
    return sentence.count(term)
def average_dl(docs):
    tot=0
    for doc in docs:
        tot+=len(doc.split(' '))
    return tot/len(docs)
def precision(relevant,total):
    return relevant/total

def tune_hyperparameters_BM25(docs,vectorizer, practices,original_paraphrased):

    true_docs={}

    excel_document = openpyxl.load_workbook('Ranked queries.xlsx')
    sheet = excel_document['Hoja1']
    all_columns = sheet.columns
    for col in all_columns:
        if col[1].value is not None:
            true_docs[col[1].value]=[]
            for row in col[2:]:
                if row.value is not None:
                    true_docs[col[1].value].append(row.value)
    
    def precision(relevant,total):
        return relevant/total
    
    def recall(relevant, tot_relevant):
        return relevant/tot_relevant
        
    def is_relevant(doc,query):
        if doc in true_docs[query]:
            return True
        else:
            return False
    def removing_original_duplicates(ranking):
        top=[]
        for doc in ranking.keys():
            original=get_original(practices[doc])
            if original not in top:
                top.append(original)
                # print(original)
                # print(ranking[doc])
        return top
    
    
    def get_original(doc):
        for k in original_paraphrased:
            if doc in original_paraphrased[k]:
                return k
    "Return the list of tokens of the document"
    def tokenizer(doc):
        tokenized_doc=nltk.word_tokenize(doc)
        return tokenized_doc
    
    "Return the stemmed version of the tokens"
    def stemmer(tokens):
        porter_stemmer=nltk.stem.PorterStemmer()
        stemmed=[]
        for word in tokens:
            stemmed.append(porter_stemmer.stem(word))
        return stemmed
    
    "Return the lemmatized version of the tokens"
    def lemmatizer(tokens):
        token_lemmatizer = WordNetLemmatizer()
        lemmatized=[]
        for word in tokens:
            lemmatized.append(token_lemmatizer.lemmatize(word))
        return lemmatized
    
    "Return tokens of the doc that are not stop words"
    def remove_stopwords(tokens):
        cleaned_practice = []
        for token in tokens:
            if token not in sw:
                cleaned_practice.append(token)
        return cleaned_practice
    def cleaning(docs, uni=True, lc=True,pun=True,num=True,sp=True,sw=True,st=True,lm=False):
        cleaned_practices = []
        all_tokens=[]
        for doc in docs:
            #Remove Unicode
            if uni:
                document_test = re.sub(r'[^\x00-\x7F]+', ' ', doc)
            # Lowercase the document
            if lc:
                document_test = document_test.lower()
            # Remove punctuations
            if pun:
                document_test = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', document_test)
            # Remove the numbers
            if num:
                document_test = re.sub(r'[0-9]', '', document_test)
            # Remove the doubled space
            if sp:
                document_test = re.sub(r'\s{2,}', ' ', document_test)
            tokens=tokenizer(document_test)
            if sw:
                tokens=remove_stopwords(tokens)
            if st:
                tokens=stemmer(tokens)
            if lm:
                tokens=lemmatizer(tokens)
            joined_doc= ' '.join(tokens)
            cleaned_practices.append(joined_doc)
            all_tokens.append(tokens)
        return cleaned_practices, all_tokens
        # return cleaned_practices
    
    queries=list(true_docs.keys())
    bs=np.arange(0.0, 1.0, 0.25)
    k1s=np.arange(0.0, 3.0, 0.2)
    b_result=[]
    k1_result=[]
    preci=[]
    R=[]
    #Acá debería recibir n= a precisiont at n TODO
    at_k=range(1,11)
    av_precision=[]
    av_recall=[]
    tot_precision=[0] * 10
    tot_recall=[0] * 10
    tot_q=[0,0,0,0,0,0,0,0,0,0]
    cants=[]
    for k1 in k1s:
        for b in bs:
            for q in queries:
                query=q+" "+expansion(q)
                query = cleaning([query])[0][0]
                print(q)
                bm25=BM25(query,docs,vectorizer, practices,b,k1)#vsm,lda,bm25
                bm25=removing_original_duplicates(bm25)
                tot_relevant=len(true_docs[q])
                if len(bm25):
                    for n in at_k:
                        if tot_relevant>=n:
                            tot_q[n-1]+=1
                            print(tot_q)
                            rlv_retrieved=0
                            bm25a=list(bm25)[:n]
                            print('ans---------------a',bm25a)
                            for r in bm25a:
                                print('ans---------------',r)
                                if is_relevant(r, q):
                                    rlv_retrieved+=1
                            print(true_docs[q])
                            precision_n=precision(rlv_retrieved,len(bm25a))
                            recall_n=recall(rlv_retrieved,tot_relevant)
                            print("Precision"+str(precision_n)+" at "+str(n))
                            print("Recall"+str(recall_n)+" at "+str(n))
                            tot_precision[n-1]+=precision_n
                            tot_recall[n-1]+=recall_n
                    cants.append("cantidad de queries validas para n="+str(tot_q))
                    print("tot_q",tot_q)
            for j in at_k:
                av_precision.append(tot_precision[j-1]/tot_q[j-1])
                av_recall.append(tot_recall[j-1]/tot_q[j-1])
            b_result.append(b)
            k1_result.append(k1)
            preci.append(sum(av_precision)/len(av_precision))
            R.append(sum(av_recall)/len(av_recall))
   
    print("CANTS-----------------",cants)
    print("b,k1"+str(b_result)+str(k1_result))
    newDocument = docx.Document()
    fig, ax = plt.subplots() 
    print("BM25")
    newDocument.add_paragraph('------Presicion BM25-------')
    newDocument.add_paragraph(str(preci))
    newDocument.add_paragraph('------Recall BM25-------')
    newDocument.add_paragraph(str(R))
    newDocument.add_paragraph('------k1s BM25-------')
    newDocument.add_paragraph(str(k1_result))
    newDocument.add_paragraph('------bs BM25-------')
    newDocument.add_paragraph(str(b_result))
    print("precision: "+str(preci))
    print("recall: "+str(R))
    ax.plot(av_recall,av_precision, color='tab:orange',marker = 'o')
    # p=0
    # for j in range(0,len(av_recall)):
    #     x=av_recall[j]
    #     y=av_precision[j]
    #     ax.annotate(str(at_k[p]),xy=(x,y),xytext =(x, y))
    #     p+=1
    ax.set_title("Precision vs Recall")
    ax.set_xlabel("Recall")
    ax.set_ylabel("Precision")
    plt.show()

    newDocument.save('PR1a10BM25.docx')
# # for i in range(1,13):
# #     doc = docx.Document("PR1a10.docx")
# #     paragraphs=doc.paragraphs
# #     resultspyr=[]#VSM,LDA,BM25
# #     for paragraph in paragraphs:
# #         if not paragraph.text.startswith("-"):
# #             print(paragraph.text)
# #             temp=paragraph.text.replace("[","").replace("]","").replace(" ","").split(",")
# #             temp=[float(i) for i in temp]
# #             resultspyr.append(temp)
# # fig, ax = plt.subplots()     
# # ax.plot(range(1,11),resultspyr[0], color='tab:purple',marker = 'o')
# # ax.plot(range(1,11),resultspyr[1], color='tab:orange',marker = 'o')

# # ax.set_title("Precision and Recall vs n (VSM)")
# # ax.set_ylabel("Precision and Recall")
# # ax.set_xlabel("n")
# # ax.set_ylim(0,1)
# # plt.legend(["Precision", "Recall"], loc ="lower right")
# # plt.show()
# # fig1, ax1 = plt.subplots()     
# # ax1.plot(range(1,11),resultspyr[2], color='tab:purple',marker = 'o')
# # ax1.plot(range(1,11),resultspyr[3], color='tab:orange',marker = 'o')

# # ax1.set_title("Precision and Recall vs n (LDA)")
# # ax1.set_ylabel("Precision and Recall")
# # ax1.set_xlabel("n")
# # ax1.set_ylim(0,1)
# # plt.legend(["Precision", "Recall"], loc ="upper right")
# # plt.show()
# # fig2, ax2 = plt.subplots()     
# # ax2.plot(range(1,11),resultspyr[4], color='tab:purple',marker = 'o')
# # ax2.plot(range(1,11),resultspyr[5],color='tab:orange',marker = 'o')

# # ax2.set_title("Precision and Recall vs n (BM25)")
# # ax2.set_ylabel("Precision and Recall")
# # ax2.set_xlabel("n")
# # ax2.set_ylim(0,1)
# # plt.legend(["Precision", "Recall"], loc ="lower right")
# # plt.show()

# # fig, ax = plt.subplots()     
# # ax.plot(resultspyr[1],resultspyr[0], color='tab:purple',marker = 'o')
# # ax.plot(resultspyr[3],resultspyr[2], color='tab:orange',marker = 'o')
# # ax.plot(resultspyr[5],resultspyr[4], color='tab:blue',marker = 'o')
# # ax.set_title("Precision vs Recall")
# # ax.set_ylabel("Precision")
# # ax.set_xlabel("Recall")
# # ax.set_ylim(0,1)
# # ax.set_xlim(0,1)
# # plt.legend(["VSM", "LDA", "BM25"], loc ="upper right")
# # plt.show()

# # fig, ax = plt.subplots()     
# # ax.plot(range(1,11),resultspyr[0], color='tab:purple',marker = 'o')
# # ax.plot(range(1,11),resultspyr[2], color='tab:orange',marker = 'o')
# # ax.plot(range(1,11),resultspyr[4], color='tab:blue',marker = 'o')
# # ax.set_title("Precision")
# # ax.set_ylabel("Precision")
# # ax.set_xlabel("n")
# # ax.set_ylim(0,1)
# # plt.legend(["VSM", "LDA", "BM25"], loc ="upper right")
# # plt.show()

# # fig, ax = plt.subplots()     
# # ax.plot(range(1,11),resultspyr[1], color='tab:purple',marker = 'o')
# # ax.plot(range(1,11),resultspyr[3], color='tab:orange',marker = 'o')
# # ax.plot(range(1,11),resultspyr[5], color='tab:blue',marker = 'o')
# # ax.set_title("Recall")
# # ax.set_ylabel("Precision")
# # ax.set_xlabel("Recall")
# # ax.set_ylim(0,1)
# # plt.legend(["VSM", "LDA", "BM25"], loc ="upper right")
# # plt.show()
# # # precision_recall()





